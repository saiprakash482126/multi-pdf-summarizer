import os
import sys
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Set
from collections import defaultdict
import re
import json
from datetime import datetime
from dataclasses import dataclass
import tempfile
import urllib.parse
import requests
from io import BytesIO

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.shared import RGBColor, Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from dotenv import load_dotenv
from urllib.parse import urlparse, parse_qs

load_dotenv()

# Define custom styles and constants
VALID_HEADING_STYLES = ['Heading 1', 'Heading 2', 'Heading 3']
VALID_BODY_STYLES = ['Normal', 'Body Text', 'List Paragraph']
VALID_TABLE_STYLES = ['Table Grid', 'Table Normal']

@dataclass
class ValidationResult:
    issue: str
    location: str
    details: str = ""
    severity: str = "Warning"  # Can be 'Error', 'Warning', or 'Info'

class DocumentValidator:
    def __init__(self, file_path: str):
        self.document_path = file_path
        self.doc: Optional[DocxDocument] = None
        self.issues: List[ValidationResult] = []
        self.styles_used: Set[str] = set()
        self.hyperlinks: List[Dict] = []
        self.bookmarks: Set[str] = set()
        self.images: List[Dict] = []
        self.tables: List[Dict] = []
        self.sections: List[Dict] = []
        self.lists: List[Dict] = []
        self.temp_file = None

    def download_sharepoint_file(self, url: str) -> Optional[bytes]:
        """Download a file from SharePoint using Microsoft Graph API and return its content in memory."""
        try:
            print(f"\n🔗 Processing SharePoint URL: {url}")
            
            # Extract file ID from different URL formats
            file_id = None
            if 'sourcedoc=' in url:
                # Extract the sourcedoc parameter and remove URL encoding
                file_id = url.split('sourcedoc=')[1].split('&')[0]
                # Remove URL encoding from the file ID
                import urllib.parse
                file_id = urllib.parse.unquote(file_id)
                # Remove any remaining curly braces
                file_id = file_id.strip('{}')
                print(f"Extracted file_id: {file_id}")
            elif '/_layouts/15/Doc.aspx' in url:
                # Try to extract from the full URL
                parsed_url = urlparse(url)
                query_params = parse_qs(parsed_url.query)
                if 'sourcedoc' in query_params:
                    file_id = query_params['sourcedoc'][0].strip('{}')
                    print(f"Extracted file_id from query params: {file_id}")
            
            if not file_id:
                # If it's a direct download URL, use the last part as file ID
                file_id = url.split('/')[-1].split('?')[0]
                print(f"Using last part of URL as file_id: {file_id}")
                
            if not file_id:
                error_msg = "Could not extract file ID from URL"
                print(f"❌ {error_msg}")
                self.issues.append(ValidationResult(
                    "Configuration Error",
                    "Invalid SharePoint URL format",
                    error_msg,
                    "Error"
                ))
                return None

            # Get authentication details
            access_token = os.getenv('ACCESS_TOKEN')
            drive_id = os.getenv('DRIVE_ID')
            
            if not all([access_token, drive_id]):
                error_msg = "Missing required environment variables (ACCESS_TOKEN and DRIVE_ID are required for SharePoint access)"
                print(f"❌ {error_msg}")
                self.issues.append(ValidationResult(
                    "Configuration Error",
                    "Missing Credentials",
                    error_msg,
                    "Error"
                ))
                return None

            print("🔑 Using provided credentials for SharePoint access")
            
            # First try to get file metadata to verify access
            headers = {
                'Authorization': f'Bearer {access_token}',
                'Accept': 'application/json'
            }
            
            # Try different endpoints with the properly formatted file ID
            endpoints = [
                f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}",
                f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/root:/{file_id}",
                f"https://graph.microsoft.com/v1.0/me/drive/items/{file_id}",
                f"https://graph.microsoft.com/v1.0/sites/root/drives/{drive_id}/items/{file_id}"
            ]
            
            file_metadata = None
            for endpoint in endpoints:
                try:
                    print(f"\n🔍 Trying endpoint: {endpoint}")
                    response = requests.get(endpoint, headers=headers, timeout=30)
                    response.raise_for_status()
                    file_metadata = response.json()
                    print("✅ Successfully retrieved file metadata")
                    break
                except requests.exceptions.RequestException as e:
                    print(f"⚠️  Failed to access {endpoint}: {str(e)}")
                    if hasattr(e, 'response') and e.response is not None:
                        print(f"Status code: {e.response.status_code}")
                        if e.response.status_code == 401:
                            print("❗ Authentication failed. The access token might be expired or invalid.")
                            print("Please check your ACCESS_TOKEN in the .env file.")
                        print(f"Response: {e.response.text[:500]}")
                    continue
            
            if not file_metadata:
                error_msg = "Failed to retrieve file metadata from any endpoint"
                print(f"❌ {error_msg}")
                return None
            
            # Try to get download URL
            download_url = file_metadata.get('@microsoft.graph.downloadUrl')
            if not download_url and 'file' in file_metadata:
                download_url = file_metadata['file'].get('@microsoft.graph.downloadUrl')
            
            if not download_url:
                # Fallback to direct download
                download_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/content"
                print(f"ℹ️  Using fallback download URL: {download_url}")
            
            # Download the file
            print(f"\n⬇️  Downloading file content...")
            response = requests.get(download_url, headers=headers, stream=True, timeout=60)
            response.raise_for_status()
            
            content = response.content
            print(f"✅ Successfully downloaded {len(content)} bytes")
            return content

        except requests.exceptions.HTTPError as e:
            error_msg = f"HTTP Error: {str(e)}"
            if hasattr(e, 'response') and e.response is not None:
                error_msg += f"\nStatus code: {e.response.status_code}"
                try:
                    error_details = e.response.json()
                    error_msg += f"\nResponse: {json.dumps(error_details, indent=2)}"
                except:
                    error_msg += f"\nResponse: {e.response.text[:1000]}"
            print(f"❌ {error_msg}")
            
        except Exception as e:
            error_msg = f"Unexpected error: {str(e)}"
            print(f"❌ {error_msg}")
            import traceback
            traceback.print_exc()
            
        self.issues.append(ValidationResult(
            "Download Failed",
            "SharePoint",
            "Failed to download file from SharePoint. Please check the URL and your permissions.",
            "Error"
        ))
        return None

    def load_document(self) -> bool:
        """Load the document from the provided path or URL."""
        if not self.document_path:
            self.issues.append(ValidationResult(
                "Configuration Error",
                "Document Path",
                "No document path provided",
                "Error"
            ))
            return False
            
        try:
            # Check file extension
            file_ext = os.path.splitext(self.document_path)[1].lower()
            if file_ext == '.pdf':
                raise ValueError("PDF files are not supported. Please provide a Word document (.docx)")
            elif file_ext not in ['.docx', '.doc']:
                raise ValueError(f"Unsupported file type: {file_ext}. Please provide a Word document (.docx)")

            # Handle SharePoint URLs
            if 'sharepoint.com' in self.document_path.lower():
                print("\n🔗 Detected SharePoint URL. Attempting to download document...")
                file_content = self.download_sharepoint_file(self.document_path)
                if not file_content:
                    return False
                    
                # Create a temporary file
                self.temp_file = os.path.join(tempfile.gettempdir(), f"temp_doc_{os.urandom(8).hex()}.docx")
                with open(self.temp_file, 'wb') as f:
                    f.write(file_content)
                
                self.doc = Document(self.temp_file)
            else:
                # Handle local file
                if not os.path.exists(self.document_path):
                    raise FileNotFoundError(f"File not found: {self.document_path}")
                    
                self.doc = Document(self.document_path)
            
            # If we got here, document was loaded successfully
            self.collect_document_data()
            return True
            
        except FileNotFoundError as e:
            error_msg = str(e)
            print(f"\n❌ File not found: {error_msg}")
            self.issues.append(ValidationResult(
                "File Not Found",
                "Document",
                error_msg,
                "Error"
            ))
            return False
            
        except ValueError as e:
            error_msg = str(e)
            print(f"\n❌ {error_msg}")
            self.issues.append(ValidationResult(
                "Unsupported File Type",
                "Document",
                error_msg,
                "Error"
            ))
            return False
            
        except Exception as e:
            error_msg = str(e)
            print(f"\n❌ Error loading document: {error_msg}")
            print("Please ensure:")
            print("1. The file exists and is accessible")
            print("2. You have the necessary permissions")
            print("3. The file is not corrupted")
            if 'sharepoint' in self.document_path.lower():
                print("4. Your SharePoint URL and credentials are correct")
            
            self.issues.append(ValidationResult(
                "Document Load Error",
                "Document",
                f"Failed to load document: {error_msg}",
                "Error"
            ))
            return False

    def cleanup(self):
        """Clean up temporary files."""
        if self.temp_file and os.path.exists(self.temp_file):
            try:
                os.remove(self.temp_file)
            except Exception:
                pass  # Ignore cleanup errors

    def collect_document_data(self) -> None:
        """Collect and store document data for validation."""
        if not self.doc:
            return
            
        # Collect styles
        self.styles_used = {style.name for style in self.doc.styles}
        
        # Collect hyperlinks
        for rel in self.doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                self.hyperlinks.append({
                    'target': rel._target,
                    'text': self.find_hyperlink_text(rel.rId)
                })
                
        # Collect tables
        for i, table in enumerate(self.doc.tables, 1):
            self.tables.append({
                'index': i,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'has_header': self.table_has_header(table),
                'style': table.style.name if table.style else None
            })
    
    def find_hyperlink_text(self, rel_id: str) -> str:
        """Find the display text of a hyperlink by its relationship ID."""
        try:
            # First try the modern approach without namespaces
            for paragraph in self.doc.paragraphs:
                for run in paragraph.runs:
                    # Check if this run has a hyperlink with the given rel_id
                    if run._element.xpath(f'.//w:hyperlink[contains(@r:id, "{rel_id}")]'):
                        return run.text.strip()
            
            # If not found, try a more comprehensive search
            for part in self.doc.part.rels.values():
                if part.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink':
                    if part.rId == rel_id:
                        # Find the hyperlink element that references this relationship
                        for paragraph in self.doc.paragraphs:
                            for run in paragraph.runs:
                                if run._element.xpath(f'.//w:hyperlink[contains(@r:id, "{rel_id}")]'):
                                    return run.text.strip()
            
            return "[Hyperlink text not found]"
            
        except Exception as e:
            print(f"Warning: Could not extract hyperlink text: {str(e)}")
            return "[Hyperlink]"
    
    def table_has_header(self, table) -> bool:
        """Check if table has a header row."""
        if not table.rows:
            return False
        first_row = table.rows[0]
        return any(cell.text.strip() for cell in first_row.cells)
    
    # ===== Validation Methods =====
    
    def run_all_checks(self) -> List[ValidationResult]:
        """Run all validation checks."""
        if not self.load_document():
            return self.issues or []
            
        # Document Properties and Metadata
        self.check_document_properties()
        self.check_document_protection()
        self.check_metadata()
        self.check_corrupt_styles()
        
        # Text and Formatting
        self.check_styles()
        self.check_text_formatting()
        self.check_hidden_text()
        self.check_highlighted_text()
        self.check_track_changes()
        self.check_manual_page_breaks()
        self.check_special_characters()
        self.check_whitespace_issues()
        self.check_language_settings()
        
        # Hyperlinks and References
        self.check_hyperlinks()
        self.check_bookmarks()
        self.check_cross_references()
        self.check_table_of_contents()
        
        # Tables
        self.check_tables()
        self.check_table_formatting()
        self.check_table_structure()
        
        # Images and Objects
        self.check_images()
        self.check_embedded_objects()
        self.check_charts()
        
        # Headers, Footers, and Page Numbers
        self.check_headers_footers()
        self.check_page_numbers()
        self.check_watermarks()
        
        # Lists and Numbering
        self.check_lists()
        self.check_numbering_consistency()
        
        # Sections and Layout
        self.check_sections()
        self.check_page_layout()
        self.check_margins()
        
        # Advanced Formatting
        self.check_fields()
        self.check_footnotes_endnotes()
        self.check_comments()
        self.check_revisions()
        
        # Content Validation
        self.check_placeholders()
        self.check_duplicate_content()
        self.check_orphan_widow_control()
        
        # Additional validation checks
        self.check_blue_text_not_linked()
        self.check_heading_hyperlink_issues()
        self.check_image_issues()
        self.check_table_issues()
        self.check_non_black_text()
        self.check_special_characters_advanced()
        self.check_orphan_captions()
        self.check_special_keywords()
        self.check_strikethrough_text()
        self.check_unused_bookmarks()
        self.check_weblinks()
        
        # New validation checks
        self.check_empty_key_properties()
        self.check_external_references()
        self.check_instructional_text()
        self.check_internal_links()
        self.check_internal_references()
        self.check_caption_labels()
        self.check_language_issues()
        self.check_narrative_formatting()
        self.check_page_layout_advanced()
        self.check_standard_styles()
        
        return self.issues or []  # Ensure we always return a list, even if empty

    # Document Properties and Metadata
    def check_document_properties(self) -> None:
        """Check core document properties."""
        if not self.doc:
            return
            
        props = self.doc.core_properties
        required_props = {
            'title': 'Document Title',
            'subject': 'Document Subject',
            'author': 'Author Name',
            'keywords': 'Document Keywords',
            'category': 'Document Category'
        }
        
        for prop, display_name in required_props.items():
            if not getattr(props, prop, None):
                self.issues.append(ValidationResult(
                    issue="Missing Document Property",
                    location="Document Properties",
                    details=f"Missing required property: {display_name}",
                    severity="Warning"
                ))

    def check_document_protection(self) -> None:
        """Check if the document has any protection settings."""
        issues = []
        try:
            # First check if the document is password protected
            if hasattr(self.doc, 'settings') and hasattr(self.doc.settings, 'protection'):
                if self.doc.settings.protection and hasattr(self.doc.settings.protection, 'write_protected'):
                    if self.doc.settings.protection.write_protected:
                        issues.append(ValidationResult(
                            "Document Protection",
                            "Document",
                            "Document is write-protected.",
                            "Warning"
                        ))
            
            # Check for restricted editing
            if hasattr(self.doc, 'settings') and hasattr(self.doc.settings, 'write_protection'):
                if self.doc.settings.write_protection and hasattr(self.doc.settings.write_protection, 'enforced'):
                    if self.doc.settings.write_protection.enforced:
                        issues.append(ValidationResult(
                            "Document Protection",
                            "Document",
                            "Document has restricted editing enabled.",
                            "Warning"
                        ))
            
            # If no protection was detected
            if not issues:
                issues.append(ValidationResult(
                    "Document Protection",
                    "Document",
                    "No document protection detected.",
                    "Info"
                ))
                
        except Exception as e:
            print(f"Warning: Could not check document protection: {str(e)}")
            issues.append(ValidationResult(
                "Document Protection",
                "Document",
                f"Could not verify document protection settings: {str(e)}",
                "Warning"
            ))
            
        self.issues.extend(issues)

    def check_metadata(self) -> None:
        """Check document metadata for sensitive information."""
        if not self.doc:
            return
            
        props = self.doc.core_properties
        metadata_fields = ['author', 'last_modified_by', 'revision', 'version']
        
        for field in metadata_fields:
            value = getattr(props, field, None)
            if value:
                self.issues.append(ValidationResult(
                    issue="Document Metadata",
                    location="Document Properties",
                    details=f"Metadata field '{field}' found: {value}",
                    severity="Info"
                ))

    def check_corrupt_styles(self) -> None:
        """Check for corrupt styles."""
        if not self.doc:
            return
        # Basic implementation - can be expanded
        pass

    # Text and Formatting
    def check_text_formatting(self) -> None:
        """Check for direct formatting instead of styles."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            for run in para.runs:
                # Check for direct formatting
                if (run.bold or run.italic or run.underline or 
                    (hasattr(run.font, 'highlight_color') and run.font.highlight_color) or
                    (run.font and run.font.name and run.font.name != 'Calibri')):
                    if para.style.name == 'Normal':
                        self.issues.append(ValidationResult(
                            issue="Direct Formatting",
                            location=f"Paragraph: {para.text[:50]}...",
                            details="Direct formatting used instead of styles",
                            severity="Warning"
                        ))
                        break

    def check_hidden_text(self) -> None:
        """Check for hidden text in the document."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            for run in para.runs:
                if hasattr(run.font, 'hidden') and run.font.hidden:
                    self.issues.append(ValidationResult(
                        issue="Hidden Text",
                        location="Document Content",
                        details="Hidden text found in document",
                        severity="Warning"
                    ))
                    return

    def check_highlighted_text(self) -> None:
        """Check for highlighted text."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            for run in para.runs:
                if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                    self.issues.append(ValidationResult(
                        issue="Highlighted Text",
                        location=f"Text: {run.text[:50]}...",
                        details=f"Text contains highlighting (color: {run.font.highlight_color})",
                        severity="Warning"
                    ))

    def check_styles(self) -> None:
        """Check for style-related issues."""
        if not self.doc:
            return
            
        # Check for corrupt styles
        for style in self.doc.styles:
            if style.name == 'Normal' and not style.font.name:
                self.issues.append(ValidationResult(
                    issue="Corrupt Style",
                    location=f"Style: {style.name}",
                    details="Normal style is corrupted",
                    severity="Error"
                ))
            
            # Check for non-standard styles
            if (style.name not in VALID_HEADING_STYLES + VALID_BODY_STYLES + VALID_TABLE_STYLES and 
                not style.name.startswith('Heading') and 
                style.type == WD_STYLE_TYPE.PARAGRAPH):
                self.issues.append(ValidationResult(
                    issue="Non-Standard Style",
                    location=f"Style: {style.name}",
                    details="Non-standard style used",
                    severity="Warning"
                ))

    def check_track_changes(self) -> None:
        """Check if track changes is enabled."""
        if not self.doc:
            return
            
        # Check for track changes in the document settings
        settings = self.doc.settings
        if hasattr(settings, 'track_revisions') and settings.track_revisions:
            self.issues.append(ValidationResult(
                issue="Track Changes Enabled",
                location="Document Settings",
                details="Track changes is enabled. Please accept/reject all changes before finalizing the document.",
                severity="Error"
            ))

    def check_manual_page_breaks(self) -> None:
        """Check for manual page breaks."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            if 'lastRenderedPageBreak' in para._element.xml or 'pageBreakBefore' in para._element.xml:
                self.issues.append(ValidationResult(
                    issue="Manual Page Break",
                    location=f"Paragraph: {para.text[:50]}...",
                    details="Manual page break detected. Use styles for page breaks instead.",
                    severity="Warning"
                ))

    def check_special_characters(self) -> None:
        """Check for special characters that might cause issues."""
        special_chars = {
            '\t': 'Tab character',
            '\r\n': 'Windows line ending',
            '\r': 'Carriage return',
            '\x0c': 'Form feed',
            '\x0b': 'Vertical tab',
            '\x1a': 'Substitute character'
        }
        
        for para in self.doc.paragraphs:
            for char, desc in special_chars.items():
                if char in para.text:
                    self.issues.append(ValidationResult(
                        issue=f"Special Character: {desc}",
                        location=f"Paragraph: {para.text[:50]}...",
                        details=f"Special character found: {char.encode('unicode_escape').decode()}",
                        severity="Warning"
                    ))

    def check_whitespace_issues(self) -> None:
        """Check for whitespace issues."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            # Check for leading/trailing spaces
            if para.text != para.text.strip():
                self.issues.append(ValidationResult(
                    issue="Extra Whitespace",
                    location=f"Paragraph: {para.text[:50]}...",
                    details="Leading or trailing whitespace found",
                    severity="Warning"
                ))
            
            # Check for multiple spaces
            if '  ' in para.text:
                self.issues.append(ValidationResult(
                    issue="Multiple Spaces",
                    location=f"Paragraph: {para.text[:50]}...",
                    details="Multiple consecutive spaces found",
                    severity="Warning"
                ))

    def check_language_settings(self) -> None:
        """Check language settings."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_bookmarks(self) -> None:
        """Check for bookmark-related issues."""
        if not self.doc:
            return
            
        # Check for broken bookmarks
        for para in self.doc.paragraphs:
            for run in para.runs:
                if hasattr(run, '_element') and run._element.xpath('.//w:bookmarkStart'):
                    bookmark_name = run._element.xpath('.//w:bookmarkStart/@w:name')
                    if bookmark_name:
                        self.bookmarks.add(bookmark_name[0])
        
        # Check for duplicate bookmarks
        if len(self.bookmarks) != len(set(self.bookmarks)):
            duplicates = [name for name in self.bookmarks if list(self.bookmarks).count(name) > 1]
            self.issues.append(ValidationResult(
                issue="Duplicate Bookmarks",
                location="Document",
                details=f"Duplicate bookmarks found: {', '.join(duplicates)}",
                severity="Error"
            ))

    def check_cross_references(self) -> None:
        """Check cross-reference issues."""
        if not self.doc:
            return
            
        # This is a simplified check - actual cross-reference validation is complex
        # and might require more sophisticated parsing
        for para in self.doc.paragraphs:
            if "see " in para.text.lower() and "page" in para.text.lower():
                if not any(run.italic for run in para.runs):
                    self.issues.append(ValidationResult(
                        issue="Possible Unlinked Cross-Reference",
                        location=f"Text: {para.text[:50]}...",
                        details="Possible manual cross-reference detected. Use Word's built-in cross-referencing.",
                        severity="Warning"
                    ))

    def check_table_of_contents(self) -> None:
        """Check table of contents."""
        if not self.doc:
            return
            
        has_toc = False
        for para in self.doc.paragraphs:
            if para.style.name.startswith('TOC'):
                has_toc = True
                # Check if TOC is up to date
                if "!" in para.text or "Error!" in para.text:
                    self.issues.append(ValidationResult(
                        issue="Table of Contents Error",
                        location="Document",
                        details="Table of Contents contains errors. Please update the TOC (right-click and select 'Update Field').",
                        severity="Error"
                    ))
                break
        
        if not has_toc and len(self.doc.paragraphs) > 10:  # Only suggest TOC for longer documents
            self.issues.append(ValidationResult(
                issue="Missing Table of Contents",
                location="Document",
                details="Consider adding a Table of Contents for better navigation.",
                severity="Info"
            ))

    def check_table_formatting(self) -> None:
        """Check table formatting issues."""
        for table in self.tables:
            # Check for tables that don't fit on the page
            if table['cols'] > 7:  # More than 7 columns might be too wide
                self.issues.append(ValidationResult(
                    issue="Wide Table",
                    location=f"Table {table['index']}",
                    details=f"Table has {table['cols']} columns which may be too wide for standard page",
                    severity="Warning"
                ))
            
            # Check for small font size in tables
            for row in self.doc.tables[table['index']-1].rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if hasattr(run.font, 'size') and run.font.size and run.font.size.pt < 9:
                                self.issues.append(ValidationResult(
                                    issue="Small Font in Table",
                                    location=f"Table {table['index']}",
                                    details=f"Font size too small: {run.font.size.pt}pt",
                                    severity="Warning"
                                ))

    def check_table_structure(self) -> None:
        """Check table structure."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_embedded_objects(self) -> None:
        """Check for embedded objects."""
        if not self.doc:
            return
            
        for rel in self.doc.part.rels.values():
            if "oleObject" in rel.reltype or "package" in rel.reltype:
                self.issues.append(ValidationResult(
                    issue="Embedded Object",
                    location="Document",
                    details=f"Embedded object found: {rel.target_ref}",
                    severity="Warning"
                ))

    def check_charts(self) -> None:
        """Check charts."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_headers_footers(self) -> None:
        """Check headers and footers."""
        if not self.doc:
            return
            
        # Check for empty headers/footers
        for section in self.doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and not header.paragraphs and not header.tables:
                    self.issues.append(ValidationResult(
                        issue="Empty Header",
                        location="Document",
                        details="Empty header found",
                        severity="Info"
                    ))
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and not footer.paragraphs and not footer.tables:
                    self.issues.append(ValidationResult(
                        issue="Empty Footer",
                        location="Document",
                        details="Empty footer found",
                        severity="Info"
                    ))

    def check_images(self) -> None:
        """Check image-related issues."""
        if not self.doc or not hasattr(self.doc, 'part') or not hasattr(self.doc.part, 'rels') or not self.doc.part.rels:
            return
        
        for rel in self.doc.part.rels.values():
            if "image" in rel.reltype:
                # Check for non-inline images
                if not rel.is_external:
                    self.issues.append(ValidationResult(
                        issue="Image Not Inline",
                        location=f"Image: {rel.target_ref}",
                        details="Image should be inline with text",
                        severity="Warning"
                    ))
                
                # Check image format
                if not rel.target_ref.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                    self.issues.append(ValidationResult(
                        issue="Incorrect Image Format",
                        location=f"Image: {rel.target_ref}",
                        details="Use standard image formats (PNG, JPG, GIF)",
                        severity="Warning"
                    ))

    def check_page_numbers(self) -> None:
        """Check page numbering."""
        if not self.doc or not hasattr(self.doc, 'sections') or not self.doc.sections:
            return
            
        has_page_numbers = False
        for section in self.doc.sections:
            if not hasattr(section, 'footer') or not hasattr(section, 'first_page_footer') or not hasattr(section, 'even_page_footer'):
                continue
                
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and hasattr(footer, 'paragraphs'):
                    for para in footer.paragraphs:
                        if hasattr(para, 'text') and ("PAGE" in para.text.upper() or "NUMPAGES" in para.text.upper()):
                            has_page_numbers = True
                            break
                    if has_page_numbers:
                        break
                if has_page_numbers:
                    break
                
        if not has_page_numbers:
            self.issues.append(ValidationResult(
                issue="Page Numbers Missing",
                location="Document Footer",
                details="Document does not appear to have page numbers in any footer",
                severity="Warning"
            ))

    def check_watermarks(self) -> None:
        """Check for watermarks."""
        if not self.doc or not hasattr(self.doc, 'sections') or not self.doc.sections:
            return
            
        # Check for watermarks in headers/footers
        for section in self.doc.sections:
            if not hasattr(section, 'header') or not hasattr(section, 'first_page_header') or not hasattr(section, 'even_page_header'):
                continue
                
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if not header or not hasattr(header, 'paragraphs'):
                    continue
                    
                for para in header.paragraphs:
                    if not hasattr(para, 'runs') or not para.runs:
                        continue
                        
                    try:
                        if any(hasattr(run, 'bold') and run.bold for run in para.runs) and \
                           any(hasattr(run, 'italic') and run.italic for run in para.runs):
                            self.issues.append(ValidationResult(
                                issue="Possible Watermark Detected",
                                location="Document Header",
                                details="A possible watermark was found in the document header. Please verify if this is intentional.",
                                severity="Info"
                            ))
                            return  # Found a watermark, no need to check further
                    except Exception as e:
                        print(f"Warning: Could not check watermark: {str(e)}")
                        continue

    def check_lists(self) -> None:
        """Check list formatting."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            if para.style.name.startswith('List') or para.style.name.startswith('List '):
                # Check for manual numbering
                if para.text and para.text[0].isdigit() and ". " in para.text[:5]:
                    self.issues.append(ValidationResult(
                        issue="Manual List Numbering",
                        location=f"List item: {para.text[:50]}...",
                        details="Manual list numbering detected. Use Word's built-in list formatting.",
                        severity="Warning"
                    ))

    def check_numbering_consistency(self) -> None:
        """Check numbering consistency."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_sections(self) -> None:
        """Check sections."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_page_layout(self) -> None:
        """Check page layout."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_margins(self) -> None:
        """Check margins."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_fields(self) -> None:
        """Check fields."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_footnotes_endnotes(self) -> None:
        """Check footnotes and endnotes."""
        if not self.doc:
            return
            
        # This is a placeholder - actual footnote/endnote checking requires more complex XML parsing
        for para in self.doc.paragraphs:
            if "[footnote" in para.text.lower() or "[endnote" in para.text.lower():
                self.issues.append(ValidationResult(
                    issue="Possible Manual Footnote/Endnote",
                    location=f"Text: {para.text[:50]}...",
                    details="Possible manual footnote/endnote detected. Use Word's built-in footnote/endnote feature.",
                    severity="Warning"
                ))

    def check_comments(self) -> None:
        """Check for unresolved comments."""
        if not self.doc:
            return
            
        # This is a simplified check - actual comment parsing requires more complex XML handling
        if 'comments' in [rel.reltype for rel in self.doc.part.rels.values()]:
            self.issues.append(ValidationResult(
                "Document Contains Comments",
                "Document",
                "Document contains comments. Please review and resolve all comments before finalizing.",
                "Info"
            ))

    def check_revisions(self) -> None:
        """Check revisions."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_placeholders(self) -> None:
        """Check placeholders."""
        if not self.doc:
            return
        # Basic implementation
        pass

    def check_duplicate_content(self) -> None:
        """Check for duplicate content."""
        if not self.doc:
            return
            
        # Simple duplicate paragraph detection
        content_map = {}
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if len(text) > 50:  # Only check longer paragraphs
                if text in content_map:
                    content_map[text] += 1
                else:
                    content_map[text] = 1
        
        for text, count in content_map.items():
            if count > 1:
                self.issues.append(ValidationResult(
                    issue="Duplicate Content",
                    location=f"Repeated paragraph: {text[:50]}...",
                    details=f"This paragraph appears {count} times in the document.",
                    severity="Warning"
                ))

    def check_orphan_widow_control(self) -> None:
        """Check for orphan/widow control issues."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            # Check for single line at top of page (orphan)
            if para.paragraph_format.keep_with_next and para.paragraph_format.keep_together:
                self.issues.append(ValidationResult(
                    issue="Orphan/Widow Control",
                    location=f"Paragraph: {para.text[:50]}...",
                    details="Orphan/widow control is enabled. This is good practice for professional documents.",
                    severity="Info"
                ))
                break

    def check_hyperlinks(self) -> None:
        """Check hyperlink-related issues."""
        if not self.doc:
            return
            
        for hyperlink in self.hyperlinks:
            if not hyperlink['target']:
                self.issues.append(ValidationResult(
                    issue="Broken Hyperlink",
                    location=f"Link text: {hyperlink['text']}",
                    details="Hyperlink has no target URL",
                    severity="Error"
                ))

    def check_tables(self) -> None:
        """Check table-related issues."""
        for table in self.tables:
            if not table['has_header']:
                self.issues.append(ValidationResult(
                    issue="Table Missing Headers",
                    location=f"Table {table['index']}",
                    details="Table should have a header row",
                    severity="Warning"
                ))

    def check_blue_text_not_linked(self) -> None:
        """Check for blue text that isn't a hyperlink."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            for run in para.runs:
                try:
                    # Check if the run has font and color properties
                    if hasattr(run, 'font') and run.font and hasattr(run.font, 'color') and run.font.color:
                        # Check if the text is blue (RGB: 0,0,255)
                        if hasattr(run.font.color, 'rgb') and run.font.color.rgb == RGBColor(0, 0, 255):
                            # Check if the run has a hyperlink
                            has_hyperlink = False
                            if hasattr(run, 'element') and hasattr(run.element, 'rPr'):
                                # Look for hyperlink elements in the run's properties
                                for child in run.element.rPr:
                                    if 'link' in child.tag:
                                        has_hyperlink = True
                                        break
                            
                            if not has_hyperlink:
                                self.issues.append(ValidationResult(
                                    issue="Blue Text Not Linked",
                                    location=f"Text: {run.text[:50]}...",
                                    details="Blue text should be a hyperlink",
                                    severity="Warning"
                                ))
                except Exception as e:
                    print(f"Warning: Could not check blue text for hyperlink: {str(e)}")
                    continue

    def check_heading_hyperlink_issues(self) -> None:
        """Check heading hyperlink issues."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            if not hasattr(para, 'style') or not hasattr(para.style, 'name') or not para.style.name.startswith('Heading'):
                continue
                
            # Check for heading with hyperlink
            has_hyperlink = False
            for run in para.runs:
                if hasattr(run, 'element') and hasattr(run.element, 'rPr'):
                    for child in run.element.rPr:
                        if 'link' in child.tag:
                            has_hyperlink = True
                            break
                if has_hyperlink:
                    break
            
            if has_hyperlink:
                # Check if heading text matches hyperlink text
                heading_text = para.text.strip()
                for run in para.runs:
                    if hasattr(run, 'element') and hasattr(run.element, 'rPr'):
                        for child in run.element.rPr:
                            if 'link' in child.tag and run.text.strip() and run.text.strip() != heading_text:
                                self.issues.append(ValidationResult(
                                    issue="Heading Hyperlink Text Mismatch",
                                    location=f"Heading: {heading_text[:50]}...",
                                    details=f"Heading text does not match hyperlink text: {run.text}",
                                    severity="Warning"
                                ))
            else:
                self.issues.append(ValidationResult(
                    issue="Heading Link Not Set",
                    location=f"Heading: {para.text[:50]}...",
                    details="Heading should have a hyperlink",
                    severity="Warning"
                ))

    def check_image_issues(self) -> None:
        """Check various image-related issues."""
        if not self.doc or not hasattr(self.doc, 'part') or not hasattr(self.doc.part, 'rels') or not self.doc.part.rels:
            return
        
        for rel in self.doc.part.rels.values():
            if "image" in rel.reltype:
                # Check for non-inline images
                if not rel.is_external:
                    self.issues.append(ValidationResult(
                        issue="Image Not Inline",
                        location=f"Image: {rel.target_ref}",
                        details="Image should be inline with text",
                        severity="Warning"
                    ))
                
                # Check image format
                if not rel.target_ref.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                    self.issues.append(ValidationResult(
                        issue="Incorrect Image Format",
                        location=f"Image: {rel.target_ref}",
                        details="Use standard image formats (PNG, JPG, GIF)",
                        severity="Warning"
                    ))

    def check_table_issues(self) -> None:
        """Check for common table-related issues in the document."""
        if not self.doc or not hasattr(self.doc, 'tables') or not self.doc.tables:
            return
            
        for i, table in enumerate(self.doc.tables, 1):
            try:
                # Check if table has a header row
                has_header = False
                if table.rows:
                    # More reliable way to check for header row
                    first_row = table.rows[0]
                    if hasattr(first_row, '_tr') and hasattr(first_row._tr, 'trPr'):
                        tr_pr = first_row._tr.trPr
                        has_header = tr_pr is not None and hasattr(tr_pr, 'tblHeader') and tr_pr.tblHeader is not None
                    
                    # Fallback: Check if first row has different formatting
                    if not has_header and len(table.rows) > 1:
                        first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
                        second_row_cells = [cell.text.strip() for cell in table.rows[1].cells]
                        # If first row has text and second row is different, it's likely a header
                        if any(first_row_cells) and first_row_cells != second_row_cells:
                            has_header = True
                
                if not has_header:
                    self.issues.append(ValidationResult(
                        "Table Header",
                        f"Table {i}",
                        "Table does not have a clearly defined header row.",
                        "Warning"
                    ))
                
                # Check for merged cells
                for row_idx, row in enumerate(table.rows, 1):
                    for cell_idx, cell in enumerate(row.cells, 1):
                        try:
                            if hasattr(cell, '_tc') and hasattr(cell._tc, 'gridSpan'):
                                if cell._tc.gridSpan and cell._tc.gridSpan.val > 1:
                                    self.issues.append(ValidationResult(
                                        "Merged Cells",
                                        f"Table {i}, Cell ({row_idx},{cell_idx})",
                                        "Merged cells can cause accessibility issues.",
                                        "Warning"
                                    ))
                        except Exception as e:
                            print(f"Warning: Could not check cell at Table {i}, Row {row_idx}, Column {cell_idx}: {str(e)}")
                
                # Check for split cells (rowspan > 1)
                for row in table.rows:
                    for cell in row.cells:
                        try:
                            if hasattr(cell, '_tc') and hasattr(cell._tc, 'vMerge'):
                                if cell._tc.vMerge is not None:
                                    self.issues.append(ValidationResult(
                                        "Split Cells",
                                        f"Table {i}",
                                        "Vertically merged cells can cause accessibility issues.",
                                        "Warning"
                                    ))
                        except Exception as e:
                            print(f"Warning: Could not check cell for vertical merge: {str(e)}")
                
                # Check table style
                if hasattr(table, '_tbl') and hasattr(table._tbl, 'tblPr'):
                    try:
                        style = table.style.name if table.style else "No Style"
                        if style not in VALID_TABLE_STYLES:
                            self.issues.append(ValidationResult(
                                "Table Style",
                                f"Table {i}",
                                f"Table uses non-standard style: {style}",
                                "Info"
                            ))
                    except Exception as e:
                        print(f"Warning: Could not check table style: {str(e)}")
                
            except Exception as e:
                print(f"Error processing table {i}: {str(e)}")
                self.issues.append(ValidationResult(
                    "Table Processing Error",
                    f"Table {i}",
                    f"Error while processing table: {str(e)}",
                    "Error"
                ))

    def check_non_black_text(self) -> None:
        """Check for non-black text."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font and run.font.color and run.font.color.rgb != RGBColor(0, 0, 0):
                    self.issues.append(ValidationResult(
                        issue="Non-Black Text",
                        location=f"Text: {run.text[:50]}...",
                        details=f"Text color is {run.font.color.rgb}",
                        severity="Warning"
                    ))

    def check_special_characters_advanced(self) -> None:
        """Check for nonbreaking hyphens and spaces."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            if "\u2011" in para.text:  # Nonbreaking hyphen
                self.issues.append(ValidationResult(
                    issue="Nonbreaking Hyphen",
                    location=f"Text: {para.text[:50]}...",
                    details="Nonbreaking hyphen detected. Use regular hyphens unless necessary.",
                    severity="Info"
                ))
            if "\u00A0" in para.text:  # Nonbreaking space
                self.issues.append(ValidationResult(
                    issue="Nonbreaking Space",
                    location=f"Text: {para.text[:50]}...",
                    details="Nonbreaking space detected. Use regular spaces unless necessary.",
                    severity="Info"
                ))

    def check_orphan_captions(self) -> None:
        """Check for orphaned captions."""
        if not self.doc:
            return
            
        for i, para in enumerate(self.doc.paragraphs):
            if para.style.name == 'Caption':
                # Check if next paragraph is not a table or figure
                if i + 1 < len(self.doc.paragraphs):
                    next_para = self.doc.paragraphs[i + 1]
                    if not any(tbl for tbl in self.doc.tables if tbl._element == next_para._p.getparent()):
                        self.issues.append(ValidationResult(
                            issue="Orphan Caption",
                            location=f"Caption: {para.text[:50]}...",
                            details="Caption is not followed by a table or figure",
                            severity="Warning"
                        ))

    def check_special_keywords(self) -> None:
        """Check for special keywords that should be avoided."""
        special_keywords = ["draft", "confidential", "for review", "todo"]
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            text = para.text.lower()
            for keyword in special_keywords:
                if keyword in text:
                    self.issues.append(ValidationResult(
                        issue="Special Keyword Used",
                        location=f"Text: {para.text[:50]}...",
                        details=f"Avoid using special keyword: {keyword}",
                        severity="Warning"
                    ))

    def check_strikethrough_text(self) -> None:
        """Check for strikethrough text."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font and run.font.strike:
                    self.issues.append(ValidationResult(
                        issue="Strikethrough Text",
                        location=f"Text: {run.text[:50]}...",
                        details="Strikethrough text should be removed or properly indicated",
                        severity="Warning"
                    ))

    def check_unused_bookmarks(self) -> None:
        """Check for unused bookmarks in the document."""
        if not self.doc or not hasattr(self.doc, 'part'):
            return
            
        try:
            # Get all bookmarks in the document
            all_bookmarks = {}
            used_bookmarks = set()
            
            # First, collect all bookmarks from the document
            if hasattr(self.doc.part, 'element') and hasattr(self.doc.part.element, 'xpath'):
                try:
                    bookmarks = self.doc.part.element.xpath('//w:bookmarkStart', namespaces=nsmap)
                    for bm in bookmarks:
                        bm_id = bm.get('{%s}id' % nsmap['w']) if hasattr(bm, 'get') else None
                        bm_name = bm.get('{%s}name' % nsmap['w']) if hasattr(bm, 'get') else None
                        if bm_id and bm_name:
                            all_bookmarks[bm_id] = bm_name
                except Exception as e:
                    print(f"Warning: Could not process bookmarks: {str(e)}")
            
            # Then find all used bookmarks in hyperlinks
            for para in self.doc.paragraphs:
                for run in para.runs:
                    try:
                        if hasattr(run, 'element') and hasattr(run.element, 'xpath'):
                            # Find hyperlinks that reference bookmarks
                            hyperlinks = run.element.xpath('.//w:hyperlink', namespaces=nsmap)
                            for hl in hyperlinks:
                                if hasattr(hl, 'get'):
                                    anchor = hl.get('{%s}anchor' % nsmap['w'])
                                    if anchor:
                                        used_bookmarks.add(anchor)
                    except Exception as e:
                        print(f"Warning: Could not process hyperlinks: {str(e)}")
                        continue
            
            # Check for unused bookmarks
            for bm_id, bm_name in all_bookmarks.items():
                if bm_name and bm_name not in used_bookmarks and not bm_name.startswith('_'):
                    self.issues.append(ValidationResult(
                        "Unused Bookmark",
                        f"Bookmark: {bm_name}",
                        "This bookmark is defined but not referenced by any hyperlink.",
                        "Info"
                    ))
            
            # Also check for broken hyperlinks (bookmark not found)
            for bm_name in used_bookmarks:
                if bm_name and bm_name not in all_bookmarks.values() and not bm_name.startswith('_'):
                    self.issues.append(ValidationResult(
                        "Broken Bookmark Reference",
                        f"Reference: {bm_name}",
                        "This hyperlink references a bookmark that doesn't exist in the document.",
                        "Warning"
                    ))
                    
        except Exception as e:
            print(f"Error checking bookmarks: {str(e)}")
            import traceback
            traceback.print_exc()
            self.issues.append(ValidationResult(
                "Bookmark Check Error",
                "Document",
                f"Could not complete bookmark validation: {str(e)}",
                "Error"
            ))

    def check_internal_links(self) -> None:
        """Check for internal links in the document."""
        if not self.doc:
            return
            
        try:
            for para in self.doc.paragraphs:
                for run in para.runs:
                    try:
                        # Check for hyperlinks in the run
                        if hasattr(run, 'element') and hasattr(run.element, 'xpath'):
                            # Find hyperlinks that might be internal
                            hyperlinks = run.element.xpath('.//*[local-name()="hyperlink"]')
                            for hl in hyperlinks:
                                if hasattr(hl, 'get'):
                                    # Get the anchor attribute without using namespaces
                                    anchor = hl.get('anchor')
                                    if anchor:
                                        # Check if this is an internal link
                                        self.issues.append(ValidationResult(
                                            "Internal Link",
                                            f"Link: {run.text[:50]}...",
                                            f"Document contains internal link to: {anchor}",
                                            "Info"
                                        ))
                    except Exception as e:
                        print(f"Warning: Could not process internal links: {str(e)}")
                        continue
                        
        except Exception as e:
            print(f"Error checking internal links: {str(e)}")
            import traceback
            traceback.print_exc()
            self.issues.append(ValidationResult(
                "Internal Link Check Error",
                "Document",
                f"Could not complete internal link validation: {str(e)}",
                "Error"
            ))

    def check_weblinks(self) -> None:
        """Check for presence of web links."""
        if not self.doc:
            return
            
        try:
            for para in self.doc.paragraphs:
                for run in para.runs:
                    try:
                        # Check for hyperlinks in the run
                        if hasattr(run, 'hyperlinks') and run.hyperlinks:
                            for hyperlink in run.hyperlinks:
                                if hasattr(hyperlink, 'address') and hyperlink.address:
                                    url = hyperlink.address
                                    if url.startswith(('http://', 'https://')):
                                        self.issues.append(ValidationResult(
                                            "Web Link",
                                            f"Link: {url}",
                                            "Document contains web links that should be verified.",
                                            "Info"
                                        ))
                        
                        # Also check for URLs in the text
                        if run.text and ('http://' in run.text or 'https://' in run.text):
                            # Simple URL detection
                            import re
                            urls = re.findall(r'https?://\S+', run.text)
                            for url in urls:
                                self.issues.append(ValidationResult(
                                    "Web Link in Text",
                                    f"URL: {url}",
                                    "Document contains web links in the text that should be properly hyperlinked.",
                                    "Warning"
                                ))
                                
                    except Exception as e:
                        print(f"Warning: Could not process run for web links: {str(e)}")
                        continue
                        
        except Exception as e:
            print(f"Error checking web links: {str(e)}")
            import traceback
            traceback.print_exc()
            self.issues.append(ValidationResult(
                "Web Link Check Error",
                "Document",
                f"Could not complete web link validation: {str(e)}",
                "Error"
            ))

    def check_empty_key_properties(self) -> None:
        """Check for empty key properties in document properties."""
        if not self.doc:
            return
            
        core_props = self.doc.core_properties
        key_properties = {
            'title': core_props.title,
            'subject': core_props.subject,
            'keywords': core_props.keywords,
            'category': core_props.category,
            'author': core_props.author
        }
        
        for prop_name, prop_value in key_properties.items():
            if not prop_value or not prop_value.strip():
                self.issues.append(ValidationResult(
                    issue="Empty Key Property",
                    location="Document Properties",
                    details=f"Document property '{prop_name}' is empty",
                    severity="Warning"
                ))

    def check_external_references(self) -> None:
        """Check formatting of external references."""
        if not self.doc:
            return
            
        for para in self.doc.paragraphs:
            # Look for common external reference patterns
            if re.search(r'\[\d+\]', para.text):
                if not para.style.name.startswith(('Bibliography', 'Reference')):
                    self.issues.append(ValidationResult(
                        issue="External Reference Formatting",
                        location=f"Text: {para.text[:50]}...",
                        details="External reference should use 'Reference' or 'Bibliography' style",
                        severity="Warning"
                    ))

    def check_instructional_text(self) -> None:
        """Check for instructional text styles."""
        if not self.doc:
            return
            
        instructional_styles = ['Instruction', 'Note', 'Tip', 'Warning']
        for para in self.doc.paragraphs:
            if any(style in para.style.name for style in instructional_styles):
                self.issues.append(ValidationResult(
                    issue="Instructional Style Used",
                    location=f"Style: {para.style.name}",
                    details=f"Instructional style '{para.style.name}' should be used consistently",
                    severity="Info"
                ))

    def check_internal_references(self) -> None:
        """Check for references to internal documents."""
        if not self.doc:
            return
            
        internal_ref_patterns = [
            r'see\s+document\s+[^\s.,;:]+',
            r'refer\s+to\s+[^\s.,;:]+',
            r'in\s+section\s+[^\s.,;:]+'
        ]
        
        for para in self.doc.paragraphs:
            text = para.text.lower()
            for pattern in internal_ref_patterns:
                if re.search(pattern, text):
                    self.issues.append(ValidationResult(
                        issue="Internal Document Reference",
                        location=f"Text: {para.text[:50]}...",
                        details="Reference to internal document detected. Ensure it's properly linked.",
                        severity="Info"
                    ))
                    break

    def check_caption_labels(self) -> None:
        """Check for invalid caption labels."""
        if not self.doc:
            return
            
        valid_labels = ['Figure', 'Table', 'Equation', 'Example', 'Listing']
        
        for para in self.doc.paragraphs:
            if para.style.name == 'Caption':
                text = para.text.strip()
                # Check if caption starts with a label
                if ':' in text:
                    label = text.split(':', 1)[0].strip()
                    if label not in valid_labels:
                        self.issues.append(ValidationResult(
                            issue="Invalid Caption Label",
                            location=f"Caption: {text[:50]}...",
                            details=f"Caption label '{label}' is not a standard label. Use one of: {', '.join(valid_labels)}",
                            severity="Warning"
                        ))

    def check_language_issues(self) -> None:
        """Check for language-related issues."""
        if not self.doc:
            return
            
        # Get document language from properties
        lang = self.doc.core_properties.language
        if not lang or 'en' not in lang.lower():
            self.issues.append(ValidationResult(
                issue="Language Setting",
                location="Document Properties",
                details=f"Document language is set to '{lang}'. Ensure correct language settings for spell checking.",
                severity="Warning"
            ))

    def check_narrative_formatting(self) -> None:
        """Check narrative text formatting consistency."""
        if not self.doc:
            return
            
        narrative_style = 'Normal'
        for para in self.doc.paragraphs:
            if not para.style.name.startswith(('Heading', 'Title', 'Subtitle', 'Quote', 'Caption')):
                if para.style.name != narrative_style:
                    self.issues.append(ValidationResult(
                        issue="Inconsistent Narrative Formatting",
                        location=f"Paragraph: {para.text[:50]}...",
                        details=f"Narrative text uses style '{para.style.name}' instead of '{narrative_style}'",
                        severity="Warning"
                    ))

    def check_page_layout_advanced(self) -> None:
        """Check advanced page layout settings."""
        if not self.doc or not hasattr(self.doc, 'sections'):
            return
            
        try:
            for i, section in enumerate(self.doc.sections, 1):
                # Check page margins
                margin_checks = [
                    ('left', section.left_margin),
                    ('right', section.right_margin),
                    ('top', section.top_margin),
                    ('bottom', section.bottom_margin),
                    ('header', section.header_distance),
                    ('footer', section.footer_distance)
                ]
                
                for name, margin in margin_checks:
                    if margin and hasattr(margin, 'inches'):
                        margin_inches = margin.inches
                        self.issues.append(ValidationResult(
                            "Page Layout",
                            f"Section {i} {name.replace('_', ' ').title()}",
                            f"{name.replace('_', ' ').title()}: {margin_inches:.2f} inches",
                            "Info"
                        ))
                
                # Check page size and orientation
                if hasattr(section, 'page_width') and hasattr(section, 'page_height'):
                    width = section.page_width.inches if hasattr(section.page_width, 'inches') else 0
                    height = section.page_height.inches if hasattr(section.page_height, 'inches') else 0
                    orientation = "Landscape" if width > height else "Portrait"
                    self.issues.append(ValidationResult(
                        "Page Layout",
                        f"Section {i} Page Size",
                        f"Page size: {width:.2f} x {height:.2f} inches ({orientation})",
                        "Info"
                    ))
                    
        except Exception as e:
            print(f"Error checking page layout: {str(e)}")
            import traceback
            traceback.print_exc()
            self.issues.append(ValidationResult(
                "Page Layout Check Error",
                "Document",
                f"Could not complete page layout validation: {str(e)}",
                "Error"
            ))

    def check_standard_styles(self) -> None:
        """Check for modified standard styles."""
        if not self.doc:
            return
            
        standard_styles = {
            'Normal': {'font_name': 'Calibri', 'font_size': 11, 'space_after': 8},
            'Heading 1': {'font_name': 'Calibri Light', 'font_size': 16, 'bold': True},
            'Heading 2': {'font_name': 'Calibri', 'font_size': 14, 'bold': True},
            'Heading 3': {'font_name': 'Calibri', 'font_size': 12, 'bold': True, 'italic': True}
        }
        
        for style_name, expected in standard_styles.items():
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                
                # Check font properties
                if hasattr(style.font, 'name') and style.font.name != expected['font_name']:
                    self.issues.append(ValidationResult(
                        issue="Modified Standard Style",
                        location=f"Style: {style_name}",
                        details=f"Font should be '{expected['font_name']}', not '{style.font.name}'",
                        severity="Warning"
                    ))
                
                # Check font size
                if hasattr(style.font, 'size') and style.font.size and \
                   style.font.size.pt != expected['font_size']:
                    self.issues.append(ValidationResult(
                        issue="Modified Standard Style",
                        location=f"Style: {style_name}",
                        details=f"Font size should be {expected['font_size']}pt, not {style.font.size.pt}pt",
                        severity="Warning"
                    ))
                
                # Check paragraph spacing
                if 'space_after' in expected and hasattr(style.paragraph_format, 'space_after'):
                    if style.paragraph_format.space_after and \
                       abs(style.paragraph_format.space_after.pt - expected['space_after']) > 0.5:
                        self.issues.append(ValidationResult(
                            issue="Modified Standard Style",
                            location=f"Style: {style_name}",
                            details=f"Spacing after should be {expected['space_after']}pt, "
                                   f"not {style.paragraph_format.space_after.pt}pt",
                            severity="Warning"
                        ))

def find_word_documents(directory: str) -> List[str]:
    """Find all Word documents in the given directory and its subdirectories."""
    word_docs = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.lower().endswith(('.doc', '.docx')):
                word_docs.append(os.path.join(root, file))
    return word_docs

def process_document(file_path: str) -> Dict[str, List[dict]]:
    """Process a single document and return its validation results."""
    print(f"\n🔍 Analyzing document: {file_path}")
    print("=" * 50)
    
    validator = DocumentValidator(file_path)
    results = defaultdict(list)
    
    try:
        # First try to load the document
        if not validator.load_document():
            error_msg = "The document could not be loaded. It may be corrupted or in an unsupported format."
            print(f"❌ {error_msg}")
            results["Error"].append({
                "file": file_path,
                "issue": "Failed to load document",
                "details": error_msg,
                "location": ""
            })
            return results
        
        # Run all validation checks
        try:
            issues = validator.run_all_checks()
            
            # Ensure issues is a list before iterating
            if not isinstance(issues, (list, tuple)):
                error_msg = f"Expected a list of issues, but got {type(issues).__name__}"
                print(f"❌ {error_msg}")
                results["Error"].append({
                    "file": file_path,
                    "issue": "Validation Error",
                    "details": error_msg,
                    "location": ""
                })
                return results
                
            # Process the issues
            issue_count = 0
            for issue in issues:
                if issue and isinstance(issue, ValidationResult):
                    results[issue.severity].append({
                        "file": file_path,
                        "issue": issue.issue,
                        "location": issue.location or "",
                        "details": issue.details or ""
                    })
                    issue_count += 1
                    
            if issue_count == 0:
                print("✅ No issues found in the document.")
                results["Info"].append({
                    "file": file_path,
                    "issue": "No Issues Found",
                    "details": "No validation issues were found in the document.",
                    "location": ""
                })
                
        except Exception as e:
            error_msg = str(e) if str(e) else "Unknown error during validation"
            print(f"❌ Validation error: {error_msg}")
            results["Error"].append({
                "file": file_path,
                "issue": "Validation Error",
                "details": f"An error occurred during validation: {error_msg}",
                "location": ""
            })
            # Print the full traceback for debugging
            import traceback
            traceback.print_exc()
                
    except Exception as e:
        error_msg = str(e) if str(e) else "An unknown error occurred"
        print(f"❌ Unexpected error: {error_msg}")
        results["Error"].append({
            "file": file_path,
            "issue": "Unexpected Error",
            "details": f"An unexpected error occurred: {error_msg}",
            "location": ""
        })
        import traceback
        traceback.print_exc()
    finally:
        try:
            validator.cleanup()
        except Exception as e:
            print(f"⚠️  Warning: Error during cleanup: {str(e)}")
    
    return results

def generate_report(validation_results: Dict[str, List[dict]], output_file: str = None) -> None:
    """Generate a report of validation results."""
    report = []
    
    # Summary counts
    total_files = len(set(result["file"] for results in validation_results.values() for result in results))
    total_issues = sum(len(issues) for issues in validation_results.values())
    
    report.append("=" * 50)
    report.append("DOCUMENT VALIDATION REPORT")
    report.append("=" * 50)
    report.append(f"\n📊 Summary:")
    report.append(f"- Total files processed: {total_files}")
    
    # Add counts by severity
    for severity in ["Error", "Warning", "Info"]:
        if severity in validation_results:
            report.append(f"- {severity}s: {len(validation_results[severity])}")
    
    # Add detailed issues by severity
    for severity in ["Error", "Warning", "Info"]:
        if severity in validation_results and validation_results[severity]:
            report.append(f"\n{'='*20} {severity.upper()}S ({len(validation_results[severity])}) {'='*20}")
            
            # Group by file
            files = defaultdict(list)
            for issue in validation_results[severity]:
                files[issue["file"]].append(issue)
            
            for file, issues in files.items():
                report.append(f"\n📄 File: {file}")
                report.append("-" * 50)
                for i, issue in enumerate(issues, 1):
                    report.append(f"\n{i}. {issue['issue']}")
                    if issue.get('location'):
                        report.append(f"   Location: {issue['location']}")
                    if issue.get('details'):
                        report.append(f"   Details: {issue['details']}")
    
    # Write to file if output_file is specified, otherwise print to console
    report_text = "\n".join(report)
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(report_text)
        print(f"\n✅ Report saved to: {output_file}")
    else:
        print("\n" + report_text)

def main():
    """Main function to run the validator."""
    if len(sys.argv) != 2 or sys.argv[1] in ['-h', '--help']:
        print("Usage: python document_validator.py <file_or_directory>")
        print("\nOptions:")
        print("  <file_or_directory>  Path to a Word document or directory containing documents")
        print("  -h, --help          Show this help message and exit")
        return 1

    input_path = sys.argv[1]
    
    # Initialize results dictionary
    all_results = defaultdict(list)
    
    # Check if input is a file or directory
    if os.path.isfile(input_path):
        # Process single file
        results = process_document(input_path)
        for severity, issues in results.items():
            all_results[severity].extend(issues)
    elif os.path.isdir(input_path):
        # Process all Word documents in directory
        word_docs = find_word_documents(input_path)
        if not word_docs:
            print(f"No Word documents found in: {input_path}")
            return 1
            
        print(f"Found {len(word_docs)} Word documents to process...")
        
        for doc_path in word_docs:
            results = process_document(doc_path)
            for severity, issues in results.items():
                all_results[severity].extend(issues)
    else:
        print(f"Error: '{input_path}' is not a valid file or directory")
        return 1
    
    # Generate and display report
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_file = f"validation_report_{timestamp}.txt"
    generate_report(all_results, report_file)
    
    # Return non-zero exit code if there are errors
    return 1 if any(severity == "Error" and issues for severity, issues in all_results.items()) else 0

if __name__ == "__main__":
    sys.exit(main())






































































