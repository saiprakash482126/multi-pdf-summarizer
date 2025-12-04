import os
import requests
import msal
import json
import nltk
from transformers import pipeline
from pathlib import Path
from dotenv import load_dotenv  
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import urllib.parse
from datetime import datetime   # ✅ FIXED: Import datetime

# Download required NLTK data
nltk.download('punkt', quiet=True)

class OneDriveSummarizer:
    def __init__(self, client_id=None, client_secret=None, tenant_id=None, access_token=None):
        if access_token:
            self.token = access_token
            self.headers = {
                "Authorization": f"Bearer {self.token}",
                "Content-Type": "application/json"
            }
        else:
            self.client_id = client_id or os.getenv('CLIENT_ID')
            self.client_secret = client_secret or os.getenv('CLIENT_SECRET')
            self.tenant_id = tenant_id or os.getenv('TENANT_ID')
            self.authority = f"https://login.microsoftonline.com/{self.tenant_id}"
            self.scope = ["https://graph.microsoft.com/.default"]
            self.token = self._get_token()
            self.headers = {
                "Authorization": f"Bearer {self.token}",
                "Content-Type": "application/json"
            }
        
        print("Loading the summarization model (this may take a minute)...")
        self.summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
        print("Model loaded successfully!")

    def _get_token(self):
        """Get access token using client credentials flow."""
        app = msal.ConfidentialClientApplication(
            self.client_id,
            authority=self.authority,
            client_credential=self.client_secret
        )
        
        result = app.acquire_token_silent(self.scope, account=None)
        if not result:
            result = app.acquire_token_for_client(scopes=self.scope)
        
        if "access_token" in result:
            return result["access_token"]
        else:
            raise Exception(f"Could not acquire token: {result.get('error_description')}")

    def list_files_in_folder(self, folder_path):
        """List all files in the specified OneDrive folder."""
        drive_id = os.getenv('DRIVE_ID')
        if not drive_id:
            raise ValueError("DRIVE_ID is not set in the .env file")
        
        folder_path = folder_path.strip('/')
        folder_path = urllib.parse.quote(folder_path)
        
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{folder_path}:/children"
        
        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            return response.json().get('value', [])
        except Exception as e:
            print(f"Error listing files: {str(e)}")
            return []

    def download_file_to_memory(self, file_id):
        """Download file content to memory and return as bytes."""
        drive_id = os.getenv('DRIVE_ID')
        if not drive_id:
            raise ValueError("DRIVE_ID is not set in the .env file")
            
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{file_id}/content"
        
        try:
            response = requests.get(url, headers=self.headers, stream=True)
            response.raise_for_status()
            return response.content
        except Exception as e:
            print(f"Error downloading file: {str(e)}")
            if hasattr(e, 'response') and e.response is not None:
                print(f"Response: {e.response.text}")
            return None

    def summarize_text(self, text, max_length=100, min_length=30):
        """Generate a structured summary of the text with bullet points."""
        if not text.strip():
            return "No content available for summarization."

        try:
            text = ' '.join(text.split()[:1000])  # Limit to first 1000 words
            chunks = [text[i:i+500] for i in range(0, len(text), 500)]
            
            key_points = []
            for chunk in chunks:
                try:
                    # ✅ FIXED: dynamic min/max length
                    word_count = len(chunk.split())
                    dynamic_min = min(min_length, max(5, word_count // 5))
                    dynamic_max = max(dynamic_min + 5, min(max_length, word_count // 2))

                    summary = self.summarizer(
                        chunk,
                        max_length=dynamic_max,
                        min_length=dynamic_min,
                        do_sample=False
                    )
                    points = summary[0]['summary_text'].split('. ')
                    key_points.extend([p.strip() for p in points if p.strip()])
                except Exception as e:
                    print(f"Error in summarization: {str(e)}")
                    continue

            if key_points:
                key_points = key_points[:5]  
                return "• " + "\n• ".join([p for p in key_points if p])
            return "• No key points could be extracted."
        except Exception as e:
            return f"• [Error in summarization: {str(e)}]"

    # rest of your process_onedrive_folder + main() remain unchanged


    def process_onedrive_folder(self, folder_path):
        """Process all files in the specified OneDrive folder and create a summary document."""
        try:
            files = self.list_files_in_folder(folder_path)
            if not files:
                print("No files found in folder.")
                return

            doc = Document()
            doc.add_heading('Document Summaries', level=1)
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Folder: {folder_path}\n", style='Intense Quote')

            for file in files:
                if 'file' not in file:
                    continue

                file_name = file['name']
                file_ext = os.path.splitext(file_name)[1].lower()

                if file_ext not in ['.txt', '.pdf', '.docx']:
                    print(f"Skipping unsupported file type: {file_name}")
                    continue

                print(f"Processing: {file_name}")
                doc.add_heading(file_name, level=2)

                try:
                    file_content = self.download_file_to_memory(file['id'])
                    if not file_content:
                        doc.add_paragraph("• [Error: Could not download file]")
                        continue

                    text = ""
                    if file_ext == '.txt':
                        text = file_content.decode('utf-8', errors='replace')
                    elif file_ext == '.pdf':
                        import PyPDF2
                        try:
                            with BytesIO(file_content) as pdf_file:
                                reader = PyPDF2.PdfReader(pdf_file)
                                text = "\n".join(
                                    [page.extract_text() for page in reader.pages if page.extract_text()]
                                )
                        except Exception as e:
                            doc.add_paragraph(f"• [Error reading PDF: {str(e)}]")
                            continue
                    elif file_ext == '.docx':
                        from docx import Document as DocxDocument
                        try:
                            docx = DocxDocument(BytesIO(file_content))
                            text = "\n".join([para.text for para in docx.paragraphs if para.text])
                        except Exception as e:
                            doc.add_paragraph(f"• [Error reading DOCX: {str(e)}]")
                            continue

                    if text.strip():
                        summary = self.summarize_text(text)
                        doc.add_paragraph("Summary:", style='Heading 3')
                        doc.add_paragraph(summary)
                    else:
                        doc.add_paragraph("• No extractable text content found.")

                except Exception as e:
                    doc.add_paragraph(f"• Error processing file: {str(e)}")

                doc.add_paragraph()

            output_filename = f"Document_Summaries_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(output_filename)

            drive_id = os.getenv('DRIVE_ID')
            upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{folder_path}/{output_filename}:/content"

            with open(output_filename, "rb") as f:
                response = requests.put(
                    upload_url,
                    headers={
                        "Authorization": f"Bearer {self.token}",
                        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    },
                    data=f
                )

            os.remove(output_filename)

            if response.status_code in (200, 201):
                print(f"\n✅ Summary document uploaded to OneDrive: {output_filename}")
            else:
                print(f"\n❌ Upload failed: {response.status_code} - {response.text}")

        except Exception as e:
            print(f"\n❌ Error: {str(e)}")


def main():
    print("=" * 60)
    print("OneDrive Document Summarizer".center(60))
    print("=" * 60)
    print()

    load_dotenv()
    access_token = os.getenv('ACCESS_TOKEN')

    try:
        if access_token:
            print("Using provided access token")
            summarizer = OneDriveSummarizer(access_token=access_token)
        else:
            print("Using client credentials flow")
            client_id = os.getenv('CLIENT_ID')
            client_secret = os.getenv('CLIENT_SECRET')
            tenant_id = os.getenv('TENANT_ID')

            if not all([client_id, client_secret, tenant_id]):
                print("Error: Missing required environment variables.")
                return

            summarizer = OneDriveSummarizer(
                client_id=client_id,
                client_secret=client_secret,
                tenant_id=tenant_id
            )

        folder_path = input("\nEnter the OneDrive folder path (e.g., Documents/Attachments): ").strip()
        if not folder_path:
            print("No folder path provided. Exiting...")
            return

        summarizer.process_onedrive_folder(folder_path)

    except Exception as e:
        print(f"\nAn error occurred: {str(e)}")


if __name__ == "__main__":
    main()